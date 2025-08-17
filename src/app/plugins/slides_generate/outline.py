# app/plugins/slides_generate/outline.py

from __future__ import annotations
from typing import List, Dict, Any
import re, json

# Prefer your repo's utilities if present
try:
    from app.tools.office_io import fetch_to_tmp, docx_to_outline as _docx_to_outline
except Exception:
    fetch_to_tmp = None
    _docx_to_outline = None

# PDF extraction fallback
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:  # pragma: no cover
    def pdf_extract_text(path: str) -> str:
        raise RuntimeError("pdfminer.six not available; install it or use your repo's doc_ops.")

# ----------------------------- Helpers ---------------------------------

def _is_slide_list(obj: Any) -> bool:
    if not isinstance(obj, list) or not obj:
        return False
    # allow list of dicts with 'title' key
    ok = 0
    for x in obj:
        if isinstance(x, dict) and "title" in x:
            ok += 1
    return ok >= max(1, len(obj) // 2)

def _looks_like_schema_stub(obj: Any) -> bool:
    """
    Detect the echoed schema example like:
    [{"title":"","bullets":["",""],"notes":"","layout_hint":"auto"}]
    """
    if not isinstance(obj, list) or len(obj) != 1:
        return False
    x = obj[0]
    if not isinstance(x, dict):
        return False
    # empty-ish slide
    if (x.get("title", "") == "" and isinstance(x.get("bullets", []), list) and len(x["bullets"]) <= 2):
        return True
    return False

def _parse_json_safe(txt: str) -> Any | None:
    try:
        return json.loads(txt)
    except Exception:
        return None

def _extract_best_json_array(raw: str) -> list | None:
    """
    Try multiple strategies and pick the *best* array:
    1) Fenced ```json ... ``` blocks (prefer the last, longest)
    2) Entire content if starts with '['
    3) Brute-force scan substrings between '[' and ']' (pick with most slides)
    """
    candidates: list[list] = []

    # 1) fenced code blocks
    for m in re.finditer(r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", raw, re.I):
        obj = _parse_json_safe(m.group(1))
        if isinstance(obj, list) and _is_slide_list(obj) and not _looks_like_schema_stub(obj):
            candidates.append(obj)

    # 2) whole body
    body = raw.strip()
    if body.startswith("[") and body.endswith("]"):
        obj = _parse_json_safe(body)
        if isinstance(obj, list) and _is_slide_list(obj) and not _looks_like_schema_stub(obj):
            candidates.append(obj)

    # 3) brute-force scan: try substrings from every '[' to every ']' after it
    # This is O(n^2) but strings are small and it's a last resort.
    opens = [m.start() for m in re.finditer(r"\[", raw)]
    closes = [m.end() for m in re.finditer(r"\]", raw)]
    for i in range(len(opens)):
        for j in range(len(closes)-1, -1, -1):
            if closes[j] <= opens[i]:
                break
            chunk = raw[opens[i]:closes[j]]
            # heuristic guard: avoid extremely tiny arrays like "[]"
            if len(chunk) < 8:
                continue
            obj = _parse_json_safe(chunk)
            if isinstance(obj, list) and _is_slide_list(obj) and not _looks_like_schema_stub(obj):
                candidates.append(obj)
                # we still continue to maybe find a larger/better one

    if not candidates:
        return None

    # Choose the candidate with the most slides; tie-breaker = latest in text (i.e., last candidate)
    candidates.sort(key=lambda arr: (len(arr), json.dumps(arr, ensure_ascii=False).__len__()))
    return candidates[-1]

# DOCX fallback using python-docx
def _docx_to_outline_fallback(docx_path: str, max_slides: int = 12) -> List[Dict[str, Any]]:
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required to parse DOCX if office_io.docx_to_outline isn't available.") from e

    doc = Document(docx_path)
    slides = []
    cur = {"title": None, "bullets": []}
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (getattr(p.style, "name", "") or "").lower()
        if "heading 1" in style or "title" in style:
            if cur["title"]:
                slides.append(cur)
            cur = {"title": text, "bullets": []}
        elif "heading" in style or p.style is None:
            cur["bullets"].append(text)
        else:
            cur["bullets"].append(text)

    if cur["title"]:
        slides.append(cur)

    if not slides:
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paras:
            title = paras[0][:80]
            bullets = paras[1:]
            slides = [{"title": title, "bullets": bullets[i:i+6]} for i in range(0, len(bullets), 6)]
            slides.insert(0, {"title": title, "bullets": []})

    return slides[:max_slides]

# ------------------------ Main outline builders ------------------------

def outline_from_prompt(mr, tr, prompt: str, lang: str = "ar", max_slides: int = 12) -> List[Dict[str, Any]]:
    SYS = (
        "أنت كاتب عروض تقديمية محترف.\n"
        "أعد *فقط* مصفوفة JSON واحدة بدون أي شرح أو نص إضافي.\n"
        "الشكل المطلوب: [{\"title\":\"\",\"bullets\":[\"\",\"\"],\"notes\":\"\",\"layout_hint\":\"auto\"}]\n"
        f"اكتب باللغة: {lang}. لا تكرر التعليمات أو المثال أعلاه.\n"
        f"عدد الشرائح: <= {max_slides}، وبحد أقصى 6 نقاط لكل شريحة.\n"
        "اجعل الشريحة 1 غلافًا بعنوان و(اختياري) سطرٍ فرعي في notes.\n"
    )
    USR = (
        f"الموضوع: {prompt}\n"
        "أنتج الآن المصفوفة النهائية داخل كتلة كود واحدة بصيغة JSON فقط."
    )
    msgs = [
        {"role": "system", "content": SYS},
        {"role": "user", "content": USR},
    ]
    raw = mr.chat(msgs)["text"]

    arr = _extract_best_json_array(raw)
    if arr is None:
        # last chance: try to parse the first fenced chunk even if it looks like a stub
        m = re.search(r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", raw, re.I)
        if m:
            try:
                arr = json.loads(m.group(1))
            except Exception:
                pass

    if not isinstance(arr, list):
        raise ValueError("Model did not return a valid JSON array of slides.")

    # Normalize basic keys & trim
    norm: List[Dict[str, Any]] = []
    for s in arr[:max_slides]:
        if not isinstance(s, dict):
            continue
        title = str(s.get("title", "")).strip()
        bullets = s.get("bullets") or []
        if isinstance(bullets, list):
            bullets = [b if isinstance(b, str) else (b.get("text","") if isinstance(b, dict) else str(b)) for b in bullets]
        else:
            bullets = [str(bullets)]
        norm.append({
            "title": title or "شريحة",
            "bullets": bullets[:6],
            "notes": s.get("notes", ""),
            "layout_hint": s.get("layout_hint", "auto"),
        })
    if not norm:
        raise ValueError("Failed to normalize slides from model output.")
    return norm

def outline_from_docx(tr, word_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
    if fetch_to_tmp is None:
        raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
    tmp = fetch_to_tmp(word_url)
    if _docx_to_outline:
        return _docx_to_outline(tmp["path"], max_slides=max_slides)
    return _docx_to_outline_fallback(tmp["path"], max_slides=max_slides)

def outline_from_pdf(tr, pdf_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
    if fetch_to_tmp is None:
        raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
    tmp = fetch_to_tmp(pdf_url)
    text = pdf_extract_text(tmp["path"])[:80_000]
    prompt = (
        "حوّل هذا النص إلى عرض شرائح منظم بمحتوى موجز ومرتب منطقيًا. "
        "اجعل الشريحة 1 غلافًا بعنوان شامل، ثم قسم المحتوى إلى شرائح ذات عناوين واضحة ونقاط مختصرة:\n\n"
        + text
    )
    class DummyMR:
        def __init__(self, mr): self.mr = mr
        def chat(self, msgs): return self.mr.chat(msgs)
    mr2 = DummyMR(tr) if hasattr(tr, "chat") else None
    if mr2 is None:
        # naive fallback without LLM
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        slides = []
        cur = {"title": "ملخص الوثيقة", "bullets": []}
        for ln in lines:
            if 0 < len(ln) < 80 and (ln.isupper() or ln.endswith(":")):
                if cur["bullets"]:
                    slides.append(cur)
                cur = {"title": ln.strip(":").strip(), "bullets": []}
            else:
                cur["bullets"].append(ln[:120])
        if cur["bullets"]:
            slides.append(cur)
        return slides[:max_slides]
    return outline_from_prompt(mr2, tr, prompt, lang="ar", max_slides=max_slides)

# # app/plugins/slides_generate/outline.py
# from __future__ import annotations

# from typing import List, Dict, Any
# import re, json

# # Prefer your repo's utilities if present
# try:
#     from app.tools.office_io import fetch_to_tmp, docx_to_outline as _docx_to_outline
# except Exception:
#     fetch_to_tmp = None
#     _docx_to_outline = None

# # PDF extraction fallback
# try:
#     from pdfminer.high_level import extract_text as pdf_extract_text
# except Exception:  # pragma: no cover
#     def pdf_extract_text(path: str) -> str:
#         raise RuntimeError("pdfminer.six not available; install it or use your repo's doc_ops.")

# # --- Helpers -----------------------------------------------------------------

# _SCHEMA_PATTERN = re.compile(r"\[\s*\{\s*\"title\"\s*:\s*\"\"", re.I)
# _JSON_FENCE_1   = re.compile(r"```json\s*(\[[\s\S]*?\])\s*```", re.I)
# _JSON_FENCE_ANY = re.compile(r"```[\w-]*\s*(\[[\s\S]*?\])\s*```", re.I)
# _JSON_ANY       = re.compile(r"(\[[\s\S]*\])", re.S)

# def _sanitize_slide(sl: dict, default_title: str) -> dict:
#     title = str(sl.get("title", "") or "").strip()

#     # If model echoed the schema array inside title, wipe it.
#     if _SCHEMA_PATTERN.search(title) or (title.startswith("[") and title.endswith("]")):
#         title = default_title

#     # Sometimes models append the schema after a title via an em dash.
#     title = re.sub(r"\s*[—-]\s*\[.*\]\s*$", "", title).strip()

#     # Coerce bullets into list[str] or list[dict{text,subs}]
#     bullets = sl.get("bullets") or []
#     norm_bullets: list[dict] = []
#     for b in bullets:
#         if isinstance(b, dict):
#             text = str(b.get("text", "")).strip()
#             subs = [str(s).strip() for s in (b.get("subs") or []) if str(s).strip()]
#         else:
#             text = str(b).strip()
#             subs = []
#         if text:
#             norm_bullets.append({"text": text, "subs": subs})

#     # Optional fields
#     notes = str(sl.get("notes") or "").strip()
#     layout_hint = sl.get("layout_hint") or "auto"

#     return {"title": title, "bullets": norm_bullets, "notes": notes, "layout_hint": layout_hint, "rtl": True}

# def _postprocess_outline(raw_list: list, prompt: str, max_slides: int) -> List[Dict[str, Any]]:
#     # Trim & sanitize
#     cleaned: list[dict] = []
#     default_cover = prompt.strip() or "العرض التقديمي"
#     for i, sl in enumerate(raw_list):
#         cleaned.append(_sanitize_slide(sl if isinstance(sl, dict) else {"title": str(sl), "bullets": []}, default_cover))

#     # Ensure a proper cover if first slide contains bullets-only or empty title
#     if not cleaned:
#         cleaned = [{"title": default_cover, "bullets": [], "notes": "", "layout_hint": "auto", "rtl": True}]
#     else:
#         if cleaned[0].get("title", "").strip() == "" and cleaned[0].get("bullets"):
#             cleaned.insert(0, {"title": default_cover, "bullets": [], "notes": "", "layout_hint": "auto", "rtl": True})
#         elif cleaned[0].get("title", "").strip() == "":
#             cleaned[0]["title"] = default_cover

#     # Cap slide & bullet counts (soft guard; quality.py also enforces)
#     cleaned = cleaned[:max_slides]
#     for sl in cleaned:
#         sl["bullets"] = sl.get("bullets", [])[:6]
#         for b in sl["bullets"]:
#             b["subs"] = b.get("subs", [])[:3]
#         sl["rtl"] = True

#     return cleaned

# # --- DOCX fallback using python-docx -----------------------------------------

# def _docx_to_outline_fallback(docx_path: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     try:
#         from docx import Document
#     except Exception as e:
#         raise RuntimeError("python-docx is required to parse DOCX if office_io.docx_to_outline isn't available.") from e

#     doc = Document(docx_path)
#     slides = []
#     cur = {"title": None, "bullets": []}
#     for p in doc.paragraphs:
#         text = (p.text or "").strip()
#         if not text:
#             continue
#         style = (getattr(p.style, "name", "") or "").lower()
#         if "heading 1" in style or "title" in style:
#             if cur["title"]:
#                 slides.append(cur)
#             cur = {"title": text, "bullets": []}
#         elif "heading" in style or p.style is None:
#             cur["bullets"].append(text)
#         else:
#             cur["bullets"].append(text)

#     if cur["title"]:
#         slides.append(cur)

#     if not slides:
#         paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
#         if paras:
#             title = paras[0][:80]
#             bullets = paras[1:]
#             slides = [{"title": title, "bullets": bullets[i:i+6]} for i in range(0, len(bullets), 6)]
#             slides.insert(0, {"title": title, "bullets": []})

#     return slides[:max_slides]

# # --- LLM paths ---------------------------------------------------------------

# def outline_from_prompt(mr, tr, prompt: str, lang: str = "ar", max_slides: int = 12) -> List[Dict[str, Any]]:
#     SYS = (
#         "أنت كاتب عروض تقديمية محترف. أعِد مخرجاتك بصيغة JSON فقط (بدون أي شرح أو نص إضافي):\n"
#         '[{"title":"","bullets":["",""],"notes":"","layout_hint":"auto"}]\n'
#         f"- عدد الشرائح: حتى {max_slides} بحد أقصى.\n"
#         "- لكل شريحة ≤ 6 نقاط، ويفضّل ≤ 7 كلمات لكل نقطة.\n"
#         "- الشريحة 1 غلاف بعنوان واضح، بدون نقاط.\n"
#         f"- اللغة: {lang}.\n"
#         "أعِد المصفوفة فقط — بدون أسطر تمهيدية أو كتل Markdown."
#     )
#     msgs = [
#         {"role": "system", "content": SYS},
#         {"role": "user", "content": f"أنشئ مخطط عرض شرائح حول: {prompt}\nالمخرجات: مصفوفة JSON فقط."},
#     ]
#     raw = mr.chat(msgs)["text"]  # may include fences or stray text

#     # Prefer fenced JSON block, else first JSON array anywhere.
#     match = _JSON_FENCE_1.search(raw) or _JSON_FENCE_ANY.search(raw) or _JSON_ANY.search(raw)
#     if not match:
#         # Defensive minimal outline
#         base = [{"title": str(prompt).strip() or "العرض", "bullets": []}]
#         return _postprocess_outline(base, prompt, max_slides)

#     txt = match.group(1)
#     try:
#         data = json.loads(txt)
#         if not isinstance(data, list):
#             raise ValueError("Top-level is not a list")
#     except Exception:
#         # Last-chance: try to isolate innermost [...]
#         inner = _JSON_ANY.search(txt)
#         data = json.loads(inner.group(1)) if inner else [{"title": str(prompt), "bullets": []}]

#     return _postprocess_outline(data, prompt, max_slides)

# def outline_from_docx(tr, word_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     if fetch_to_tmp is None:
#         raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
#     tmp = fetch_to_tmp(word_url)
#     if _docx_to_outline:
#         return _postprocess_outline(_docx_to_outline(tmp["path"], max_slides=max_slides), "مستند وورد", max_slides)
#     return _postprocess_outline(_docx_to_outline_fallback(tmp["path"], max_slides=max_slides), "مستند وورد", max_slides)

# def outline_from_pdf(tr, pdf_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     if fetch_to_tmp is None:
#         raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
#     tmp = fetch_to_tmp(pdf_url)
#     text = pdf_extract_text(tmp["path"])[:80_000]
#     prompt = (
#         "حوّل هذا النص إلى عرض شرائح منظم بمحتوى موجز ومرتب منطقيًا. "
#         "اجعل الشريحة 1 غلافًا بعنوان شامل، ثم قسّم المحتوى إلى شرائح ذات عناوين واضحة ونقاط مختصرة:\n\n"
#         + text
#     )
#     # Re-use the LLM path directly
#     return outline_from_prompt(mr=type("MR", (), {"chat": lambda _, msgs: tr.chat(msgs)}) if hasattr(tr, "chat") else None,
#                                tr=tr, prompt=prompt, lang="ar", max_slides=max_slides)


# # app/plugins/slides_generate/outline.py
# from __future__ import annotations

# from typing import List, Dict, Any
# import re, json

# # Prefer your repo's utilities if present
# try:
#     from app.tools.office_io import fetch_to_tmp, docx_to_outline as _docx_to_outline
# except Exception:
#     fetch_to_tmp = None
#     _docx_to_outline = None

# # PDF extraction fallback
# try:
#     from pdfminer.high_level import extract_text as pdf_extract_text
# except Exception:  # pragma: no cover
#     def pdf_extract_text(path: str) -> str:
#         raise RuntimeError("pdfminer.six not available; install it or use your repo's doc_ops.")

# # DOCX fallback using python-docx
# def _docx_to_outline_fallback(docx_path: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     try:
#         from docx import Document
#     except Exception as e:
#         raise RuntimeError("python-docx is required to parse DOCX if office_io.docx_to_outline isn't available.") from e

#     doc = Document(docx_path)
#     slides = []
#     cur = {"title": None, "bullets": []}
#     for p in doc.paragraphs:
#         text = (p.text or "").strip()
#         if not text:
#             continue
#         style = (getattr(p.style, "name", "") or "").lower()
#         if "heading 1" in style or "title" in style:
#             # Start new slide
#             if cur["title"]:
#                 slides.append(cur)
#             cur = {"title": text, "bullets": []}
#         elif "heading" in style or p.style is None:
#             # lower heading → bullet
#             cur["bullets"].append(text)
#         else:
#             cur["bullets"].append(text)

#     if cur["title"]:
#         slides.append(cur)

#     # fallback if no headings at all: chunk paragraphs by size
#     if not slides:
#         paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
#         if paras:
#             title = paras[0][:80]
#             bullets = paras[1:]
#             slides = [{"title": title, "bullets": bullets[i:i+6]} for i in range(0, len(bullets), 6)]
#             slides.insert(0, {"title": title, "bullets": []})

#     return slides[:max_slides]


# def outline_from_prompt(mr, tr, prompt: str, lang: str = "ar", max_slides: int = 12) -> List[Dict[str, Any]]:
#     SYS = (
#         "You are a senior presentation writer. Produce a JSON array of slides with this schema:\n"
#         '[{"title": "...", "bullets": ["...", "..."], "notes": "", "layout_hint": "auto"}]\n'
#         "Rules: 8–14 slides unless specified, ≤6 bullets/slide, aim ≤7 words/bullet (soft).\n"
#         "Use language: " + lang + ". Avoid duplication across slides.\n"
#         "If the topic implies a cover slide, make slide 1 a cover with just title and optional subtitle.\n"
#     )
#     msgs = [
#         {"role": "system", "content": SYS},
#         {"role": "user", "content": f"Create an outline for: {prompt}\nSlides: {max_slides}"},
#     ]
#     raw = mr.chat(msgs)["text"]
#     # Extract the first JSON array from the response
#     m = re.search(r"\[.*\]", raw, re.S)
#     if not m:
#         # try fenced code blocks
#         m = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", raw, re.S)
#     if not m:
#         raise ValueError("Model did not return JSON array for slides.")
#     try:
#         data = json.loads(m.group(1) if m.lastindex else m.group(0))
#     except Exception as e:
#         raise ValueError("Failed to parse JSON outline from model output.") from e
#     return data[:max_slides]


# def outline_from_docx(tr, word_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     if fetch_to_tmp is None:
#         raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
#     tmp = fetch_to_tmp(word_url)
#     if _docx_to_outline:
#         return _docx_to_outline(tmp["path"], max_slides=max_slides)
#     return _docx_to_outline_fallback(tmp["path"], max_slides=max_slides)


# def outline_from_pdf(tr, pdf_url: str, max_slides: int = 12) -> List[Dict[str, Any]]:
#     if fetch_to_tmp is None:
#         raise RuntimeError("fetch_to_tmp not available. Wire office_io.fetch_to_tmp in your repo.")
#     tmp = fetch_to_tmp(pdf_url)
#     text = pdf_extract_text(tmp["path"])[:80_000]
#     # Chunk giant text and ask model to structure it
#     prompt = (
#         "حوّل هذا النص إلى عرض شرائح منظم بمحتوى موجز ومرتب منطقيًا. "
#         "اجعل الشريحة 1 غلافًا بعنوان شامل، ثم قسم المحتوى إلى شرائح ذات عناوين واضحة ونقاط مختصرة:\n\n"
#         + text
#     )
#     # Re-use the LLM pathway
#     class DummyMR:
#         def __init__(self, mr): self.mr = mr
#         def chat(self, msgs): return self.mr.chat(msgs)
#     mr = DummyMR(tr) if hasattr(tr, "chat") else None
#     if mr is None:
#         # fall back: naive split by headings-like lines
#         lines = [l.strip() for l in text.splitlines() if l.strip()]
#         slides = []
#         cur = {"title": "ملخص الوثيقة", "bullets": []}
#         for ln in lines:
#             if len(ln) > 0 and (ln.isupper() or ln.endswith(":")) and len(ln) < 80:
#                 if cur["bullets"]:
#                     slides.append(cur)
#                 cur = {"title": ln.strip(":").strip(), "bullets": []}
#             else:
#                 cur["bullets"].append(ln[:120])
#         if cur["bullets"]:
#             slides.append(cur)
#         return slides[:max_slides]
#     # With model
#     return outline_from_prompt(mr, tr, prompt, lang="ar", max_slides=max_slides)
